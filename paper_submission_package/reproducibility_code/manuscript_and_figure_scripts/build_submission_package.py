from __future__ import annotations

import csv
import hashlib
import shutil
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(r"E:\Demo of GVIM\deer-flow-mainnew")
SOURCE_DOC = ROOT / "GVIM2.0.docx"
SOURCE_SI = ROOT / "Supporting_Information" / "GVIM2.0_Supporting_Information.docx"
PACKAGE = ROOT / "GVIM2.0_Submission_Package_2026-06-22"

THREADS = [
    {
        "demo": "ESOL solubility regression",
        "user_id": "4d0cfbcc-f7f4-4a8e-be86-96856f93447a",
        "thread_id": "a515769f-4313-4160-b578-614bef7b409f",
        "frontend_scope": "Full front-end workflow: feature generation, scaffold-aware 5-fold CV, model selection, refit, and unlabeled-test prediction",
        "endpoint_provenance": "Front-end CV metrics; held-out test metrics computed post hoc after freezing predictions and joining withheld public ESOL labels",
        "verified_endpoint": "CV selected RandomForest: MAE 0.6429, RMSE 0.8541, R2 0.8554; post-hoc held-out test: MAE 0.603, RMSE 0.741, R2 0.739",
    },
    {
        "demo": "Matbench steels fold 0",
        "user_id": "4d0cfbcc-f7f4-4a8e-be86-96856f93447a",
        "thread_id": "b4442f6a-a613-4598-9b55-448eaaa33c7c",
        "frontend_scope": "Front-end workflow for official Matbench fold 0 only: internal 5-fold CV on 249 training rows, model selection, refit, and prediction of 63 unlabeled fold-0 test rows",
        "endpoint_provenance": "Internal CV metrics are front-end outputs; fold-0 test metrics were computed post hoc after freezing predictions and joining withheld public labels",
        "verified_endpoint": "Internal CV selected GradientBoosting: MAE 87.788, RMSE 123.779, R2 0.8192; post-hoc fold-0 test: MAE 111.444 MPa, RMSE 164.320 MPa, R2 0.684",
    },
    {
        "demo": "Publisher-JATS reaction-table extraction",
        "user_id": "f37c0628-e18f-4d93-9d56-fa16c1efbe50",
        "thread_id": "9dc67557-7d82-476d-80e0-381b7af1e76f",
        "frontend_scope": "Front-end PDF parsing and reaction-table artifact generation",
        "endpoint_provenance": "Post-hoc deterministic comparison with publisher JATS XML gold table",
        "verified_endpoint": "Exact precision/recall/F1 0.981/0.981/0.981; row accuracy 0.905 over 105 cells",
    },
    {
        "demo": "MassBank MS/MS candidate retrieval",
        "user_id": "4214b5a2-6e1a-42da-9c2a-697cfe80f9de",
        "thread_id": "abe44a97-e5a9-4ba3-bbf3-6206ab9d7c12",
        "frontend_scope": "Front-end spectrum search and ranked-candidate output",
        "endpoint_provenance": "Post-hoc deterministic ranking evaluation using public MassBank InChIKey identities",
        "verified_endpoint": "Top-1 0.800; Top-3 1.000; Top-5 1.000; MRR 0.900; 5 queries",
    },
    {
        "demo": "ChEMU chemical entity extraction",
        "user_id": "a167ae3d-c222-47d6-a040-0f7b4e4e9ba1",
        "thread_id": "dbb54347-66bc-4244-8177-bdf441383454",
        "frontend_scope": "Front-end entity extraction and span-prediction output",
        "endpoint_provenance": "Post-hoc exact and relaxed span evaluation against the uploaded ChEMU gold sample",
        "verified_endpoint": "Exact precision/recall/F1 0.931/0.728/0.817; relaxed F1 0.841; 92 gold spans",
    },
    {
        "demo": "BACE retrospective active discovery",
        "user_id": "0d253518-3b38-4d96-aba9-4675f48f3fca",
        "thread_id": "fff9cae7-3467-496d-af05-0585c99fd993",
        "frontend_scope": "Full front-end retrospective active-discovery simulation with 20 seeds",
        "endpoint_provenance": "Metrics emitted directly by the front-end run from public retrospective labels",
        "verified_endpoint": "Random Recall@150 0.109; greedy-surrogate 0.333; UCB 0.363; UCB EF@150 3.664",
    },
    {
        "demo": "Matbench experimental band-gap modelling and discovery",
        "user_id": "93f3f850-ffb2-4e84-a113-04eb856fd34d",
        "thread_id": "a6600dc1-8a72-46d1-97c0-d6ce0020c7f6",
        "frontend_scope": "Full front-end 5-fold KFold regression on 4,604 rows plus a three-seed retrospective active-discovery simulation",
        "endpoint_provenance": "Regression and discovery metrics emitted directly by the front-end run",
        "verified_endpoint": "5-fold MAE 0.4513 +/- 0.0170 eV, RMSE 0.8049 +/- 0.0456 eV, R2 0.6854 +/- 0.0445; greedy Recall@150 0.2626",
    },
    {
        "demo": "BACE1 temporal external validation",
        "user_id": "2473c8a4-9275-4d73-8c2f-92a8b6472310",
        "thread_id": "9ae4e85f-542b-4f94-9294-81189cf220be",
        "frontend_scope": "Full two-stage front-end workflow: pre-2016 training/model selection and post-2018 external evaluation",
        "endpoint_provenance": "Metrics emitted directly by the front-end workflow after label-isolated Stage 1",
        "verified_endpoint": "External MAE 0.8215, RMSE 1.0278, R2 0.3687; Recall@10% 0.2609; EF10 2.6054; n 1,598",
    },
]

POSTHOC_FILES = [
    ROOT / "research-demos/results/fast_ml_showcase/frontend_esol_scaffold_aware.json",
    ROOT / "research-demos/results/fast_ml_showcase/frontend_steels_single_run_recheck.json",
    ROOT / "research-demos/results/mineru_reaction_table/frontend_9dc67557_evaluation.json",
    ROOT / "research-demos/results/msms_retrieval/frontend_abe44a97_evaluation.json",
    ROOT / "research-demos/results/chemu/frontend_dbb54347_evaluation.json",
]

EXCLUDED_FILES = [
    ROOT / "research-demos/results/matbench_steels_5fold/independent_evaluation.json",
]


def thread_path(item: dict[str, str]) -> Path:
    return (
        ROOT
        / "deer-flow-main/.deer-flow/users"
        / item["user_id"]
        / "threads"
        / item["thread_id"]
    )


def delete_row(table, row) -> None:
    table._tbl.remove(row._tr)


def set_paragraph(paragraph, text: str) -> None:
    style = paragraph.style
    alignment = paragraph.alignment
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    paragraph.style = style
    paragraph.alignment = alignment


def replace_paragraph_start(document: Document, start: str, new_text: str) -> None:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(start)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {start!r}, found {len(matches)}")
    set_paragraph(matches[0], new_text)


def correct_main_document(source: Path, destination: Path) -> None:
    doc = Document(source)

    # Table 2: retain only the actually front-end-executed steels fold-0 workflow.
    for row in list(doc.tables[1].rows):
        if row.cells[0].text.strip() == "Matbench steels 5-fold":
            delete_row(doc.tables[1], row)
    doc.tables[1].cell(2, 0).text = "Matbench steels fold 0"
    doc.tables[1].cell(2, 1).text = "Matbench v0.1 official fold 0; frozen predictions scored post hoc"

    # Table 3: remove a duplicate alias. The archived implementation has random,
    # greedy_surrogate, and ucb_surrogate policies only.
    for row in list(doc.tables[2].rows):
        if row.cells[0].text.strip() == "Active search ET":
            delete_row(doc.tables[2], row)
    doc.tables[2].cell(2, 0).text = "Greedy-surrogate ET"
    doc.tables[2].cell(3, 0).text = "UCB-surrogate ET"

    # Table 4: exclude the local independent official-five-fold steels rerun.
    for row in list(doc.tables[3].rows):
        if "official 5-fold" in row.cells[0].text.lower():
            delete_row(doc.tables[3], row)
    doc.tables[3].cell(2, 0).text = "Matbench steels fold 0"
    doc.tables[3].cell(2, 1).text = "Matbench v0.1 official fold 0; front-end predictions with post-hoc withheld-label scoring"

    replace_paragraph_start(
        doc,
        "Figure 1.",
        "Figure 1. Experimental design and evaluation workflow of GVIM 2.0. The study combines a paired 400-task public benchmark with five front-end-initiated task-native demonstrations, two retrospective fixed-budget discovery cases, and one two-stage temporal external-validation case. Endpoints are based on native answer keys, public labels, publisher-defined gold data, or deterministic task-family metrics. No expert subjective grading or prospective wet-lab validation is claimed.",
    )
    replace_paragraph_start(
        doc,
        "To test whether GVIM can execute complete scientific workflows",
        "To test whether GVIM can execute complete scientific workflows rather than only answer questions, we audited front-end thread records spanning chemistry, materials science, literature extraction, spectral retrieval, and chemical information extraction. For ESOL and Matbench steels fold 0, the front end generated frozen predictions from unlabeled test inputs; held-out metrics were calculated only afterward by joining those predictions to withheld public labels. The other demonstrations were scored from archived front-end artifacts using public or publisher-defined references and standard task-family metrics: MAE, RMSE, and R虏 for regression; precision, recall, and F1 for extraction; and top-k accuracy and mean reciprocal rank for retrieval.",
    )
    replace_paragraph_start(
        doc,
        "Table 2.",
        "Table 2. Verified front-end demonstrations and objective endpoint metrics. Held-out ESOL and Matbench steels fold-0 metrics were computed post hoc from frozen front-end predictions; they were not available to the agent during execution.",
    )
    replace_paragraph_start(
        doc,
        "The evaluation workflows show complementary aspects",
        "The verified workflows show complementary aspects of the system. ESOL and Matbench steels fold 0 assess lightweight molecular and materials property-prediction workflows [36-38]. The steels thread does not constitute an official Matbench five-fold front-end evaluation: only official fold 0 was uploaded, while a shuffled five-fold split of its 249-row training partition was used internally for model selection. Reaction-table extraction uses publisher XML as the gold standard, MS/MS retrieval uses public MassBank InChIKey identities, and ChEMU uses exact and relaxed span scoring [39,40]. ChEMU recall remains lower than precision and is presented as evidence of task coverage rather than state-of-the-art named-entity recognition.",
    )
    replace_paragraph_start(
        doc,
        "Figure 3.",
        "Figure 3. Task-native chemistry and materials demonstrations initiated through the GVIM front end. (a) MoleculeNet ESOL: frozen predictions for 254 unlabeled test molecules were scored afterward against withheld public labels (MAE = 0.603, RMSE = 0.741, R虏 = 0.739). (b) Matbench steels official fold 0: frozen predictions for 63 unlabeled test alloys were scored afterward against withheld labels (MAE = 111.444 MPa, RMSE = 164.320 MPa, R虏 = 0.684). This is not a complete official five-fold Matbench front-end evaluation. (c) Publisher-JATS reaction-table extraction over 105 cells. (d) MassBank candidate retrieval over five queries. (e) ChEMU entity extraction over 92 gold spans. (f) Coverage of the five workflow types. Metrics are task-native and are not pooled into a custom aggregate score.",
    )
    replace_paragraph_start(
        doc,
        "Table 4.",
        "Table 4. Verified GVIM front-end artifacts and objective metrics. Direct-output metrics are distinguished from post-hoc scoring of frozen front-end predictions against withheld public gold labels.",
    )
    replace_paragraph_start(
        doc,
        "Figure 4.",
        "Figure 4. Retrospective BACE active discovery under a fixed query budget using public pIC50 labels. High-activity targets were the top 5% of 1,513 molecules (n = 79). Each run began with 30 random queries and selected batches of 10 up to 150 queries using an ExtraTrees surrogate; 20 random seeds were used. (a) Recall@150, (b) enrichment factor@150, and (c) best recovered pIC50 for random, greedy-surrogate and upper-confidence-bound policies. Bars show means, points show individual seeds, and error bars show 卤1 SD; the dashed line in (b) denotes random-prevalence enrichment (EF = 1). (d) Mean recall improvement over random with 95% confidence intervals from 10,000 paired per-seed bootstrap resamples. Public labels were revealed only when queried and for final retrospective evaluation; no wet-lab validation is claimed.",
    )
    replace_paragraph_start(
        doc,
        "The real GVIM output shows that the lightweight composition-only model",
        "The archived GVIM output shows that the lightweight composition-only model achieved a shuffled five-fold KFold MAE of 0.4513 eV and R虏 of 0.6854 on all 4,604 records. This was executed by the front-end thread and is distinct from the Matbench steels fold-0 case. In the retrospective active-discovery simulation, greedy_surrogate increased Recall@150 from 0.0332 to 0.2626. Because only three seeds were run, the reported paired bootstrap interval of 0.2208-0.2381 is treated as exploratory and descriptive rather than as strong inferential evidence.",
    )

    # Remove any surviving official-five-fold steels statements from body or tables.
    for paragraph in doc.paragraphs:
        if "92.584" in paragraph.text or "137.135" in paragraph.text:
            raise RuntimeError(f"Uncorrected independent steels metric in paragraph: {paragraph.text}")
    for table in doc.tables:
        for row in table.rows:
            text = " | ".join(cell.text for cell in row.cells)
            if "92.584" in text or "137.135" in text:
                raise RuntimeError(f"Uncorrected independent steels metric in table: {text}")

    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def correct_si_document(source: Path, destination: Path) -> None:
    doc = Document(source)
    replace_paragraph_start(
        doc,
        "Each workflow was initiated from the GVIM user-facing workflow",
        "Each reported workflow is linked to an archived GVIM front-end thread and is scored against a public or publisher-defined reference. Matbench steels is limited to official fold 0: the front end used internal shuffled five-fold cross-validation on the 249-row fold-0 training partition, refitted the selected model, and predicted 63 unlabeled fold-0 test records. The separate local official-five-fold rerun is excluded from the manuscript evidence because it was not executed through the GVIM front end.",
    )
    replace_paragraph_start(
        doc,
        "ESOL used the MoleculeNet ESOL held-out test set",
        "For ESOL, the front end performed scaffold-aware five-fold cross-validation, selected RandomForest, and generated frozen predictions for 254 unlabeled test molecules; post-hoc scoring against withheld public labels gave MAE = 0.603, RMSE = 0.741 and R虏 = 0.739. For Matbench steels, the front end processed official fold 0 only and generated frozen predictions for 63 unlabeled test alloys; post-hoc scoring gave MAE = 111.444 MPa, RMSE = 164.320 MPa and R虏 = 0.684. No complete official Matbench steels five-fold front-end result is claimed.",
    )
    for table in doc.tables:
        for row in list(table.rows):
            if "official 5-fold" in row.cells[0].text.lower():
                delete_row(table, row)
    for paragraph in doc.paragraphs:
        if "92.584" in paragraph.text or "137.135" in paragraph.text:
            raise RuntimeError(f"Uncorrected independent steels metric in SI paragraph: {paragraph.text}")
    destination.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destination)


def replace_docx_media(docx_path: Path, replacements: dict[str, Path]) -> None:
    """Replace embedded raster bytes while preserving Word layout and sizing."""
    temporary = docx_path.with_suffix(".media-replaced.docx")
    if temporary.exists():
        temporary.unlink()
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED
    ) as destination:
        for item in source.infolist():
            data = source.read(item.filename)
            if item.filename in replacements:
                replacement = replacements[item.filename]
                if not replacement.exists():
                    raise FileNotFoundError(replacement)
                data = replacement.read_bytes()
            destination.writestr(item, data)
    try:
        temporary.replace(docx_path)
    except PermissionError:
        docx_path.unlink()
        temporary.replace(docx_path)


def copy_evidence() -> None:
    if PACKAGE.exists():
        shutil.rmtree(PACKAGE)
    (PACKAGE / "manuscript").mkdir(parents=True)
    (PACKAGE / "supporting_information").mkdir(parents=True)
    (PACKAGE / "frontend_thread_records").mkdir(parents=True)
    (PACKAGE / "posthoc_gold_scoring").mkdir(parents=True)
    (PACKAGE / "excluded_non_frontend").mkdir(parents=True)

    correct_main_document(SOURCE_DOC, PACKAGE / "manuscript/GVIM2.0.docx")
    correct_si_document(SOURCE_SI, PACKAGE / "supporting_information/GVIM2.0 Support Information.docx")

    # Only replace figures with confirmed data/provenance defects. Their aspect
    # ratios match the existing Word placements, so pagination is preserved.
    corrected_figures = {
        "word/media/image4.png": ROOT / "manuscript_assets/figures/Figure2_demo_landscape.png",
        "word/media/image5.png": ROOT / "manuscript_assets/figures/Fig4_experiment5_bace_active_discovery.png",
        "word/media/image6.png": ROOT / "manuscript_assets/figures/Fig5_matbench_bandgap_main_case.png",
    }
    replace_docx_media(PACKAGE / "manuscript/GVIM2.0.docx", corrected_figures)

    figure_dir = PACKAGE / "publication_figures_600dpi"
    figure_dir.mkdir(parents=True, exist_ok=True)
    publication_figures = [
        ("Figure3_task_native_demos", ROOT / "manuscript_assets/figures/Figure2_demo_landscape"),
        ("Figure4_BACE_active_discovery", ROOT / "manuscript_assets/figures/Fig4_experiment5_bace_active_discovery"),
        ("Figure5_matbench_bandgap", ROOT / "manuscript_assets/figures/Fig5_matbench_bandgap_main_case"),
    ]
    for output_stem, source_stem in publication_figures:
        for extension in (".png", ".svg"):
            source = source_stem.with_suffix(extension)
            if not source.exists():
                raise FileNotFoundError(source)
            shutil.copy2(source, figure_dir / f"{output_stem}{extension}")

    manifest_rows = []
    for item in THREADS:
        source = thread_path(item)
        if not source.exists():
            raise FileNotFoundError(source)
        destination = PACKAGE / "frontend_thread_records" / item["thread_id"]
        shutil.copytree(source, destination)
        manifest_rows.append({**item, "archived_path": str(destination.relative_to(PACKAGE))})

    with (PACKAGE / "frontend_evidence_manifest.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(manifest_rows[0]))
        writer.writeheader()
        writer.writerows(manifest_rows)

    for source in POSTHOC_FILES:
        if not source.exists():
            raise FileNotFoundError(source)
        target = PACKAGE / "posthoc_gold_scoring" / source.parent.name / source.name
        target.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, target)

    for source in EXCLUDED_FILES:
        if source.exists():
            target = PACKAGE / "excluded_non_frontend" / source.parent.name / source.name
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(source, target)

    readme = """# GVIM 2.0 evidence package

This package applies a strict claim boundary: a workflow is described as a GVIM
front-end run only when its DeerFlow thread directory contains the uploaded
inputs, generated workspace code, and resulting output artifacts.

## Critical Matbench steels correction

Thread b4442f6a-a613-4598-9b55-448eaaa33c7c contains official Matbench steels
fold 0 only (249 training rows and 63 unlabeled test rows). Its five-fold KFold
operation is internal model selection on the fold-0 training partition. It is
not a complete official Matbench five-fold evaluation. The separate local
official-five-fold result is retained under excluded_non_frontend for audit
transparency and is excluded from manuscript claims about GVIM front-end runs.

## Metric provenance

Some front-end workflows produced metrics directly. ESOL and Matbench steels
fold 0 received unlabeled test files, so the front end produced frozen
predictions without test metrics. Their held-out metrics were calculated post
hoc by joining those frozen predictions to withheld public labels. This scoring
step is identified explicitly in the manifest and manuscript.

No expert subjective grading or prospective wet-lab validation is claimed.
"""
    (PACKAGE / "README.md").write_text(readme, encoding="utf-8")

    excluded_readme = """The files in this directory are retained only for audit transparency.
They were not produced by a complete GVIM front-end workflow and must not be
reported as front-end experimental evidence. In particular, the Matbench
steels official-five-fold evaluation was a separate local independent rerun.
"""
    (PACKAGE / "excluded_non_frontend/README.txt").write_text(excluded_readme, encoding="utf-8")

    hash_rows = []
    for path in sorted(PACKAGE.rglob("*")):
        if path.is_file() and path.name != "SHA256_MANIFEST.csv":
            digest = hashlib.sha256(path.read_bytes()).hexdigest()
            hash_rows.append((str(path.relative_to(PACKAGE)), path.stat().st_size, digest))
    with (PACKAGE / "SHA256_MANIFEST.csv").open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["relative_path", "size_bytes", "sha256"])
        writer.writerows(hash_rows)


if __name__ == "__main__":
    copy_evidence()
    print(PACKAGE)

