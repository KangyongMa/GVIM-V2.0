from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"E:\Demo of GVIM\deer-flow-mainnew")
SI_DIR = ROOT / "Supporting_Information"
OUTPUT = SI_DIR / "GVIM2.0_Supporting_Information.docx"
RESULTS = ROOT / "research-demos" / "results"
SOURCE = ROOT / "manuscript_assets" / "source_data"
TEMPORAL = RESULTS / "frontend_bace_temporal_9ae4e85f"

INK = "1F2937"
MUTED = "5F6D7C"
HEADER_FILL = "E8EEF5"
ALT_FILL = "F7F9FC"
ACCENT = "1F4D78"


def set_run_font(run, size: float = 10.5, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(INK)


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell(cell, value: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(1.5)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.0
    set_run_font(paragraph.add_run(str(value)), size=8.2, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, amount in (("top", 70), ("bottom", 70), ("start", 100), ("end", 100)):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(amount))
        element.set(qn("w:type"), "dxa")
        margins.append(element)


def set_table_geometry(table, widths_in: list[float]) -> None:
    total_dxa = round(sum(widths_in) * 1440)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = tbl_pr.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(total_dxa))
    table_width.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            dxa = round(width * 1440)
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")


def add_table(document: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(caption), size=9, italic=True)

    table = document.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[index], HEADER_FILL)
    for row_index, values in enumerate(rows):
        added_row = table.add_row()
        prevent_row_split(added_row)
        cells = added_row.cells
        for index, value in enumerate(values):
            set_cell(cells[index], str(value))
            if row_index % 2 == 1:
                shade(cells[index], ALT_FILL)
    set_table_geometry(table, widths)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.15
    set_run_font(paragraph.add_run(text))


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run_font(paragraph.add_run(text))


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def skill_frontmatter(path: Path) -> dict[str, str]:
    text = path.read_text(encoding="utf-8", errors="replace")
    if not text.startswith("---"):
        return {"name": path.parent.name, "description": "", "license": ""}
    block = text.split("---", 2)[1]

    def scalar(key: str) -> str:
        match = re.search(rf"(?m)^{re.escape(key)}:\s*(.*)$", block)
        if not match:
            return ""
        value = match.group(1).strip()
        if value in {">", ">-", "|", "|-"} or not value:
            lines = []
            remainder = block[match.end() :].splitlines()
            for line in remainder:
                if line and not line[0].isspace():
                    break
                if line.strip():
                    lines.append(line.strip())
            value = " ".join(lines)
        return value.strip().strip("'\"")

    return {"name": scalar("name") or path.parent.name, "description": scalar("description"), "license": scalar("license")}


def compact_description(value: str, limit: int = 180) -> str:
    value = re.sub(r"\\n", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    if len(value) <= limit:
        return value
    boundary = value.rfind(" ", 0, limit - 1)
    if boundary < 80:
        boundary = limit - 1
    return value[:boundary].rstrip(" ,;:") + "…"


def load_skill_catalog() -> list[list[str]]:
    project = ROOT / "deer-flow-main"
    skills_root = project / "skills"
    extensions = json.loads((project / "extensions_config.json").read_text(encoding="utf-8"))
    configured = extensions.get("skills", {})
    rows = []
    for path in sorted(skills_root.rglob("SKILL.md")):
        metadata = skill_frontmatter(path)
        relative = path.relative_to(project).as_posix()
        category = relative.split("/")[1]
        state = configured.get(metadata["name"], {}).get("enabled")
        config_label = "explicitly enabled" if state is True else "explicitly disabled" if state is False else "catalogued"
        declared_source = metadata["license"] or ("local custom Skill" if category == "custom" else "repository Skill definition")
        rows.append(
            [
                metadata["name"],
                category,
                config_label,
                compact_description(metadata["description"]),
                declared_source,
            ]
        )
    if len(rows) != 147:
        raise ValueError(f"Expected 147 Skill definitions, found {len(rows)}")
    return rows


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in (
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 10, 4),
        ("Heading 3", 11, 8, 3),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.add_run("GVIM 2.0 | Supporting Information"), size=8.5, italic=True)
    header.runs[-1].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("S"), size=8.5)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def build() -> None:
    benchmark = pd.read_csv(SOURCE / "Fig2_benchmark_source_data.csv")
    paired = pd.read_csv(SOURCE / "Fig2_paired_statistics.csv").iloc[0]
    api_models = pd.read_csv(SOURCE / "API_only_multimodel_400_source_data.csv")
    workflow_audit = pd.read_csv(SOURCE / "real_demo_artifact_audit.csv")
    bace = json.loads((RESULTS / "frontend_bace_fff9cae7" / "metrics.json").read_text(encoding="utf-8"))
    bandgap = json.loads((RESULTS / "frontend_fast_features_a6600dc1" / "metrics.json").read_text(encoding="utf-8"))
    bandgap_folds = pd.DataFrame(bandgap["regression_metrics"]["per_fold_scores"])
    temporal_metrics = json.loads((TEMPORAL / "metrics.json").read_text(encoding="utf-8"))
    temporal_manifest = json.loads((TEMPORAL / "prediction_manifest.json").read_text(encoding="utf-8"))
    temporal_audit = json.loads((TEMPORAL / "local_import_audit.json").read_text(encoding="utf-8"))
    temporal_cv = pd.read_csv(TEMPORAL / "cv_results.csv")
    temporal_external = pd.read_csv(TEMPORAL / "external_predictions_with_gold.csv")

    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(title.add_run("Supporting Information"), size=17, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    set_run_font(
        subtitle.add_run("GVIM 2.0: A Multi-Agent Research System for Chemistry and Materials Science"),
        size=13,
        bold=True,
    )
    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(author.add_run("Kangyong Ma"), size=10.5)
    affiliation = document.add_paragraph()
    affiliation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(
        affiliation.add_run(
            "College of Physics and Electronic Information Engineering, Zhejiang Normal University, Jinhua 321000, China"
        ),
        size=9.5,
        italic=True,
    )
    document.add_page_break()

    add_heading(document, "Contents", 1)
    contents = [
        "S1. Scope, evidence levels and claim boundaries",
        "S2. Public 400-task benchmark",
        "S3. Task-native chemistry and materials evaluation workflows",
        "S4. Retrospective BACE fixed-budget active discovery",
        "S5. Matbench experimental band-gap case",
        "S6. BACE1 temporal external validation",
        "S7. Reproducibility, software and artifact provenance",
        "S8. Metric definitions",
        "S9. Limitations",
        "S10. Supporting references",
    ]
    for item in contents:
        add_body(document, item)

    add_heading(document, "S1. Scope, evidence levels and claim boundaries", 1)
    add_body(
        document,
        "This Supporting Information documents only evaluations supported by archived GVIM artifacts or public benchmark source data. The study contains three evidence levels: public question-answer benchmarks, task-native evaluation workflows with externally defined gold standards, and retrospective discovery or temporal-validation studies using public labels. No expert subjective score is used as a primary endpoint, and no result is presented as prospective wet-lab validation.",
    )
    add_table(
        document,
        "Table S1. Evidence classes and permitted claims.",
        ["Evidence class", "Primary endpoint", "Permitted interpretation"],
        [
            ["Public benchmark", "Native answer-key accuracy", "Knowledge/reasoning performance under benchmark rules"],
            ["Task-native workflow", "MAE, RMSE, R², precision, recall, F1, top-k, MRR", "Ability to execute an auditable data-to-result workflow"],
            ["Retrospective discovery", "Recall@k, EF@k, prediction error", "Prioritization within a public labelled library"],
            ["Temporal external validation", "External MAE/RMSE/R² and top-decile recovery", "Generalization to later non-overlapping public compounds"],
        ],
        [1.45, 2.15, 2.90],
    )

    add_heading(document, "S2. Public 400-task benchmark", 1)
    add_body(
        document,
        "The suite contains 250 ChemBench, 97 MaScQA and 53 ChemLLMBench questions. Each source was scored by its native answer key or task-specific rule. Because these rules differ, the 400-question total is descriptive; per-benchmark outcomes are the primary interpretable units. The same-base comparison used the identical 400 questions for GVIM and the DeepSeek API-only setting.",
    )
    add_table(
        document,
        "Table S2. Benchmark composition and same-base results.",
        ["Benchmark", "n", "GVIM", "API-only", "GVIM (%)", "API-only (%)"],
        [[r.benchmark, int(r.n), f"{int(r.gvim_correct)}/{int(r.n)}", f"{int(r.api_correct)}/{int(r.n)}", f"{r.gvim_percent:.2f}", f"{r.api_percent:.2f}"] for r in benchmark.itertuples(index=False)],
        [1.35, 0.45, 1.00, 1.00, 1.10, 1.10],
    )
    add_body(
        document,
        f"The paired contingency table contained {int(paired.both_correct)} both-correct, {int(paired.gvim_only_correct)} GVIM-only, {int(paired.api_only_correct)} API-only and {int(paired.both_wrong)} both-wrong outcomes. The paired difference was {paired.delta_percentage_points:.1f} percentage points, with a bootstrap 95% CI of {paired.bootstrap_ci_low_pp:.1f}–{paired.bootstrap_ci_high_pp:.2f} percentage points and exact two-sided McNemar p = 3.40 × 10⁻⁶.",
    )
    add_table(
        document,
        "Table S3. Additional complete API-only reference runs on the 400-task suite.",
        ["Model", "Correct", "Accuracy (%)", "Run status"],
        [[str(r.model), f"{int(r.correct)}/{int(r.n)}", f"{r.accuracy_percent:.2f}", str(r.status)] for r in api_models.itertuples(index=False)],
        [3.05, 0.85, 1.10, 1.50],
    )

    add_heading(document, "S3. Task-native chemistry and materials evaluation workflows", 1)
    add_body(
        document,
        "Each workflow was initiated from the GVIM user-facing workflow or reconstructed from its archived output artifacts and was scored against a public or publisher-defined reference. Matbench steels fold 0 represents the front-end demonstration; the official five-fold result is retained as a separate local independent evaluation and is not described as a full front-end run.",
    )
    add_table(
        document,
        "Table S4. Artifact-supported task-native evaluations.",
        ["Workflow", "Gold/dataset", "Objective result"],
        [[r.demo, r.gold_or_dataset, r.objective_metrics.replace("+/-", "±").replace("R2", "R²")] for r in workflow_audit.itertuples(index=False)],
        [1.65, 2.00, 2.85],
    )
    add_heading(document, "S3.1 Regression workflows", 2)
    add_body(
        document,
        "ESOL used the MoleculeNet ESOL held-out test set (n = 254) and reported MAE = 0.603, RMSE = 0.741 and R² = 0.739. The Matbench steels front-end fold-0 demonstration used n = 63 test records and reported MAE = 111.444 MPa, RMSE = 164.320 MPa and R² = 0.684. The separate official five-fold evaluation reported MAE = 92.584 ± 12.313 MPa, RMSE = 137.135 ± 23.155 MPa and R² = 0.781 ± 0.087.",
    )
    add_heading(document, "S3.2 Extraction and retrieval workflows", 2)
    add_body(
        document,
        "Reaction-table extraction used the publisher JATS XML table as gold and evaluated 105 cells (21 rows). Exact precision, recall and F1 were each 0.981 and row accuracy was 0.905. The MassBank workflow ranked candidate molecular identities for five queries using public InChIKey labels, giving top-1 accuracy 0.800, top-3 and top-5 accuracy 1.000, and MRR 0.900. ChEMU Task 1 sample v3 contained 92 gold spans; exact precision/recall/F1 were 0.931/0.728/0.817 and relaxed F1 was 0.841.",
    )

    add_heading(document, "S4. Retrospective BACE fixed-budget active discovery", 1)
    params = bace["experiment_parameters"]
    add_body(
        document,
        f"The public library contained {params['n_molecules']:,} molecules. High-activity targets were the top 5% by pIC50 (n = {params['n_high_activity']}). Each of {params['n_seeds']} seeds began with {params['initial_random_queries']} random observations and queried batches of {params['batch_size']} to a total budget of {params['total_budget']}. The surrogate was {params['surrogate_model']}. Unqueried labels were not available to model fitting or acquisition and were used only as the retrospective oracle and for final scoring.",
    )
    add_table(
        document,
        "Table S5. BACE fixed-budget endpoint metrics (mean ± SD across 20 seeds).",
        ["Policy", "Recall@150", "Hit rate@150", "EF@150", "Best pIC50@150"],
        [[name.replace("_", " "), f"{values['recall_at_150']['mean']:.3f} ± {values['recall_at_150']['std']:.3f}", f"{values['hit_rate_at_150']['mean']:.3f} ± {values['hit_rate_at_150']['std']:.3f}", f"{values['enrichment_factor_at_150']['mean']:.3f} ± {values['enrichment_factor_at_150']['std']:.3f}", f"{values['best_pIC50_at_150']['mean']:.3f} ± {values['best_pIC50_at_150']['std']:.3f}"] for name, values in bace["aggregated_metrics"].items()],
        [1.55, 1.25, 1.25, 1.20, 1.25],
    )
    active_ci = bace["bootstrap_95_ci_active_minus_random_recall"]
    ucb_ci = bace["bootstrap_95_ci_ucb_surrogate_minus_random_recall"]
    add_body(
        document,
        f"Paired bootstrap intervals used 10,000 resamples of per-seed endpoint differences. Active-search minus random Recall@150 had a 95% CI of {active_ci['lower']:.3f}–{active_ci['upper']:.3f}; UCB minus random had a 95% CI of {ucb_ci['lower']:.3f}–{ucb_ci['upper']:.3f}.",
    )

    add_heading(document, "S5. Matbench experimental band-gap case", 1)
    add_body(
        document,
        f"The archived front-end case used {bandgap['dataset']} (n = {bandgap['n_samples']:,}) with {bandgap['n_features']} precomputed composition descriptors. {bandgap['regression_model']} was evaluated by {bandgap['cv_method']}. The retrospective active-discovery component defined the top 5% experimental-gap materials as targets and used three seeds, 30 initial random observations, batches of 10 and a total budget of 150.",
    )
    add_table(
        document,
        "Table S6. Per-fold Matbench experimental band-gap regression metrics.",
        ["Fold", "MAE (eV)", "RMSE (eV)", "R²"],
        [[int(r.fold), f"{r.MAE:.4f}", f"{r.RMSE:.4f}", f"{r.R2:.4f}"] for r in bandgap_folds.itertuples(index=False)],
        [1.00, 1.75, 1.75, 2.00],
    )
    aggregate = bandgap["regression_metrics"]
    active = bandgap["active_discovery_results"]
    add_body(
        document,
        f"Across folds, MAE was {aggregate['MAE_mean']:.4f} ± {aggregate['MAE_std']:.4f} eV, RMSE was {aggregate['RMSE_mean']:.4f} ± {aggregate['RMSE_std']:.4f} eV, and R² was {aggregate['R2_mean']:.4f} ± {aggregate['R2_std']:.4f}. At 150 queries, random and greedy-surrogate Recall@150 were {active['random']['recall_mean']:.4f} and {active['greedy_surrogate']['recall_mean']:.4f}, respectively. The paired recall improvement was {active['recall_diff_greedy_minus_random']['mean']:.4f} with bootstrap 95% CI {active['recall_diff_greedy_minus_random']['ci95_lower']:.4f}–{active['recall_diff_greedy_minus_random']['ci95_upper']:.4f}.",
    )

    add_heading(document, "S6. BACE1 temporal external validation", 1)
    add_heading(document, "S6.1 Data split and phase gate", 2)
    add_body(
        document,
        "ChEMBL target CHEMBL4822 records were filtered to standardized BACE1 IC50 measurements and aggregated by compound. Stage 1 used 5,296 compounds associated with documents published through 2015. Records from 2016–2017 were excluded as a buffer. The external set contained 1,598 compounds associated with documents from 2018 onward after exact molecule overlap removal by InChIKey; 1,275 compounds were scaffold-novel. The external gold file was withheld until Stage 1 had written predictions and a prediction manifest.",
    )
    add_table(
        document,
        "Table S7. Stage-1 scaffold-grouped cross-validation metrics.",
        ["Model", "Fold", "MAE", "RMSE", "R²"],
        [[r.model, r.fold, f"{r.MAE:.4f}", f"{r.RMSE:.4f}", f"{r.R2:.4f}"] for r in temporal_cv.itertuples(index=False)],
        [1.45, 0.75, 1.35, 1.35, 1.60],
    )
    add_body(
        document,
        f"RandomForest was selected by mean validation MAE and refit on all historical compounds using random seed {temporal_manifest['random_seed']}. The model used {temporal_manifest['features']} and 100 trees with n_jobs = 1. Stage-1 runtime was {temporal_manifest['runtime_seconds']:.2f} s.",
    )

    add_heading(document, "S6.2 Independent metric recomputation", 2)
    regression_rows = temporal_metrics["regression_metrics"]
    add_table(
        document,
        "Table S8. Post-2018 external regression performance.",
        ["Model", "MAE", "RMSE", "R²", "Pearson r"],
        [[r["model"], f"{r['MAE']:.4f}", f"{r['RMSE']:.4f}", f"{r['R2']:.4f}", f"{r['Pearson_r']:.4f}"] for r in regression_rows],
        [2.00, 1.10, 1.10, 1.10, 1.20],
    )
    add_table(
        document,
        "Table S9. Prespecified top-decile recovery at a 10% selection budget.",
        ["Subset", "n", "k", "Actives", "Hits", "Recall", "EF"],
        [[r["subset"].replace("scaffold_novel_0", "known scaffold").replace("scaffold_novel_1", "novel scaffold").replace("all_external", "all external"), r["n"], r["k"], r["n_actives"], r["hits_in_top_k"], f"{r['recall_at_k']:.4f}", f"{r['enrichment_factor']:.4f}"] for r in temporal_metrics["ranking_metrics"]],
        [1.40, 0.65, 0.60, 0.80, 0.65, 1.15, 1.25],
    )
    add_body(
        document,
        "Independent local recomputation from external_predictions_with_gold.csv reproduced the selected-model values: MAE = 0.8214935, RMSE = 1.0278495 and R² = 0.3687357. The file contained 1,598 unique compound_id values, no missing values and a zero maximum absolute difference between Stage-1 selected predictions and the predictions carried into the Stage-2 merged file.",
    )
    add_heading(document, "S6.3 Frozen-prediction integrity audit", 2)
    add_body(
        document,
        f"The Stage-1 manifest records SHA-256 {temporal_manifest['external_predictions_csv_sha256']}. The HTTP artifact download used LF line endings and had raw SHA-256 {temporal_audit['raw_download_sha256']}. Converting only line endings to CRLF, without changing CSV fields or values, yielded SHA-256 {temporal_audit['crlf_canonical_sha256']}, which exactly matches the manifest. This newline-sensitive distinction is retained here rather than reporting raw-byte identity. Semantic prediction identity was additionally verified by a one-to-one compound_id merge and zero maximum absolute prediction difference.",
    )

    add_heading(document, "S7. Reproducibility, software and artifact provenance", 1)
    add_heading(document, "S7.1 Explicitly enabled chemistry and materials Skills", 2)
    add_body(
        document,
        "The repository contained 147 discoverable SKILL.md files at package finalization. Table S10 reports only the six chemistry/materials Skills explicitly marked enabled in extensions_config.json. A Skill is an instruction and tool-routing contract rather than an independent scientific algorithm; citations therefore refer to the underlying software or database used for the corresponding scientific operation. The complete repository inventory is supplied as source_data/skills_inventory.csv.",
    )
    add_table(
        document,
        "Table S10. Core domain Skills explicitly configured as enabled and their scientific foundations.",
        ["Skill", "Verified role and boundary", "Skill file under skills/public/", "Scientific foundation"],
        [
            ["chemistry-studio-ketcher", "Molecule/reaction canvas preparation; not route validation", "chemistry-studio-ketcher/SKILL.md", "Ketcher [S9]; RDKit [S10]"],
            ["chemistry-structure-resolution", "Name/identifier resolution before downstream analysis", "chemistry-structure-resolution/SKILL.md", "PubChem [S14]; RDKit [S10]"],
            ["rdkit-molecular-analysis", "Descriptors, similarity, standardization, scaffolds, reaction QC and conformers", "rdkit-molecular-analysis/SKILL.md", "RDKit [S10]"],
            ["materials-core", "Formula/structure, XRD and precursor calculations; not phase-purity proof", "materials-core/SKILL.md", "pymatgen [S11]"],
            ["materials-evidence-project", "Database-backed stability, band-gap and formation-energy evidence", "materials-evidence-project/SKILL.md", "Materials Project [S12]; OPTIMADE [S13]; PubChem [S14]"],
            ["science-upload-autopilot", "Routes uploaded scientific files to deterministic tools and reports evidence gaps", "science-upload-autopilot/SKILL.md", "This work; delegates to [S10-S14]"],
        ],
        [1.35, 2.35, 1.55, 1.25],
    )

    add_heading(document, "S7.2 Complete Skill catalogue", 2)
    add_body(
        document,
        "Table S11 lists every SKILL.md definition discovered under skills/public and skills/custom. The configuration label distinguishes the six explicit enablement records from catalogue entries without an explicit override. Catalogue presence demonstrates an available instruction module, not experimental use, successful execution, preinstallation of every optional dependency, or independent validation of the capability. References for the six chemistry/materials Skills are given in Table S10; other entries retain their declared license or repository provenance rather than being assigned an unsupported scholarly citation.",
    )
    add_table(
        document,
        "Table S11. Complete repository Skill catalogue at package finalization (n = 147).",
        ["Skill", "Class", "Configuration", "Declared capability", "Declared source/license"],
        load_skill_catalog(),
        [1.25, 0.55, 0.85, 2.75, 1.10],
    )

    add_heading(document, "S7.3 Software versions and artifact provenance", 2)
    versions = temporal_manifest["package_versions"]
    add_table(
        document,
        "Table S12. Software recorded by the BACE1 temporal Stage-1 manifest.",
        ["Component", "Version"],
        [[key, value] for key, value in versions.items()],
        [2.00, 4.50],
    )
    artifact_paths = [
        RESULTS / "fast_ml_showcase" / "frontend_esol_scaffold_aware.json",
        RESULTS / "fast_ml_showcase" / "frontend_steels_single_run_recheck.json",
        RESULTS / "matbench_steels_5fold" / "independent_evaluation.json",
        RESULTS / "mineru_reaction_table" / "frontend_9dc67557_evaluation.json",
        RESULTS / "msms_retrieval" / "frontend_abe44a97_evaluation.json",
        RESULTS / "chemu" / "frontend_dbb54347_evaluation.json",
        RESULTS / "frontend_bace_fff9cae7" / "metrics.json",
        RESULTS / "frontend_fast_features_a6600dc1" / "metrics.json",
        TEMPORAL / "metrics.json",
        TEMPORAL / "prediction_manifest.json",
    ]
    add_table(
        document,
        "Table S13. Key local artifacts and SHA-256 identifiers.",
        ["Artifact", "SHA-256"],
        [[str(path.relative_to(ROOT)).replace("\\", "/"), sha256(path)] for path in artifact_paths],
        [4.35, 2.15],
    )
    add_body(
        document,
        "The temporal-validation front-end thread identifier was 9ae4e85f-542b-4f94-9294-81189cf220be. Other explicitly archived front-end identifiers include 9dc67557-7d82-476d-80e0-381b7af1e76f (reaction-table extraction), abe44a97-e5a9-4ba3-bbf3-6206ab9d7c12 (MS/MS retrieval), dbb54347-66bc-4244-8177-bdf441383454 (ChEMU), fff9cae7-3467-496d-af05-0585c99fd993 (BACE active discovery), and a6600dc1-8a72-46d1-97c0-d6ce0020c7f6 (Matbench band-gap case). Thread identifiers are reported only where directly recorded in the archived workflow history.",
    )

    add_heading(document, "S8. Metric definitions", 1)
    definitions = [
        "MAE = n⁻¹ Σᵢ |yᵢ − ŷᵢ|.",
        "RMSE = [n⁻¹ Σᵢ (yᵢ − ŷᵢ)²]¹ᐟ².",
        "R² = 1 − Σᵢ(yᵢ − ŷᵢ)² / Σᵢ(yᵢ − ȳ)².",
        "Precision = TP/(TP + FP), recall = TP/(TP + FN), and F1 is their harmonic mean.",
        "Top-k accuracy is the fraction of queries for which the gold identity appears among the first k ranked candidates; MRR is the mean reciprocal rank of the gold candidate.",
        "Recall@k is the fraction of prespecified target compounds recovered within k queries or selections.",
        "EF@k = (target fraction among selected compounds)/(target fraction in the full evaluation library).",
    ]
    for definition in definitions:
        add_bullet(document, definition)

    add_heading(document, "S9. Limitations", 1)
    for limitation in [
        "All discovery cases are retrospective and use public labels; no new compound or material was experimentally validated.",
        "The task-native extraction and retrieval demonstrations are small, fixed evaluations intended to demonstrate workflow execution rather than establish state-of-the-art task models.",
        "The 400-task aggregate combines benchmarks with different native scoring rules and must remain descriptive.",
        "The BACE1 temporal model retains non-random enrichment for scaffold-novel compounds but shows higher prediction error and systematic underprediction of some highly active novel chemotypes.",
        "Complete agent-baseline and component-ablation experiments remain necessary before attributing all gains to a specific GVIM module.",
    ]:
        add_bullet(document, limitation)

    add_heading(document, "S10. Supporting references", 1)
    references = [
        "S1. Mirza, A. et al. A framework for evaluating the chemical knowledge and reasoning abilities of large language models against the expertise of chemists. Nat. Chem. 2025, 17, 1027–1034.",
        "S2. Zaki, M.; Krishnan, N. A. MaScQA: investigating materials science knowledge of large language models. Digit. Discov. 2024, 3, 313–327.",
        "S3. Wu, Z. et al. MoleculeNet: a benchmark for molecular machine learning. Chem. Sci. 2018, 9, 513–530.",
        "S4. Dunn, A. et al. Benchmarking materials property prediction methods: the Matbench test set and Automatminer reference algorithm. npj Comput. Mater. 2020, 6, 138.",
        "S5. Delaney, J. S. ESOL: estimating aqueous solubility directly from molecular structure. J. Chem. Inf. Comput. Sci. 2004, 44, 1000–1005.",
        "S6. Horai, H. et al. MassBank: a public repository for sharing mass spectral data for life sciences. J. Mass Spectrom. 2010, 45, 703–714.",
        "S7. Nguyen, D. Q. et al. ChEMU: named entity recognition and event extraction of chemical reactions from patents. ECIR 2020, 572–579.",
        "S8. Gaulton, A. et al. ChEMBL: a large-scale bioactivity database for drug discovery. Nucleic Acids Res. 2012, 40, D1100–D1107.",
        "S9. Karulin, B.; Kozhevnikov, M. Ketcher: web-based chemical structure editor. J. Cheminform. 2011, 3 (Suppl. 1), P3. https://doi.org/10.1186/1758-2946-3-S1-P3.",
        "S10. Landrum, G. RDKit: Open-source cheminformatics software. https://www.rdkit.org (accessed 21 June 2026).",
        "S11. Ong, S. P. et al. Python Materials Genomics (pymatgen): a robust, open-source Python library for materials analysis. Comput. Mater. Sci. 2013, 68, 314–319. https://doi.org/10.1016/j.commatsci.2012.10.028.",
        "S12. Jain, A. et al. Commentary: The Materials Project: a materials genome approach to accelerating materials innovation. APL Mater. 2013, 1, 011002. https://doi.org/10.1063/1.4812323.",
        "S13. Andersen, C. W. et al. OPTIMADE, an API for exchanging materials data. Sci. Data 2021, 8, 217. https://doi.org/10.1038/s41597-021-00974-z.",
        "S14. Kim, S. et al. PubChem 2023 update. Nucleic Acids Res. 2023, 51, D1373–D1380. https://doi.org/10.1093/nar/gkac956.",
    ]
    for reference in references:
        add_body(document, reference)

    document.core_properties.title = "GVIM 2.0 Supporting Information"
    document.core_properties.subject = "Verified benchmark, workflow and temporal-validation protocols and results"
    document.core_properties.author = "Kangyong Ma"
    SI_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
